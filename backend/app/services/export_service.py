"""
Assembles all approved sections into a single, properly formatted .docx file.
Uses python-docx (runs at request-time in the user's own backend instance,
not at chat-time).
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.config import settings


def build_docx(topic: str, sections: list[dict], project_id: str) -> Path:
    """
    sections: list of {title, content} in order, only approved sections.
    Returns path to the generated .docx file.
    """
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(topic)
    run.bold = True
    run.font.size = Pt(24)
    doc.add_paragraph()  # spacer
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Research Report").italics = True
    doc.add_page_break()

    for section in sections:
        heading = doc.add_heading(section["title"], level=1)
        content = _strip_repeated_heading(section["content"].strip(), section["title"])

        if section["title"].strip().lower() == "references":
            # Each line of references as its own paragraph (already numbered by the LLM)
            for line in content.split("\n"):
                line = line.strip()
                if line:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(6)
        else:
            for para_text in content.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    p = doc.add_paragraph(para_text)
                    p.paragraph_format.space_after = Pt(8)
                    p.paragraph_format.line_spacing = 1.15

    safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in topic)[:60].strip() or "report"
    out_path = settings.export_dir / f"{project_id}_{safe_name}.docx"
    doc.save(out_path)
    return out_path


def _strip_repeated_heading(content: str, title: str) -> str:
    """
    The writer agent often restates the section title as its first line.
    build_docx already adds that heading, so drop the duplicate line.
    """
    lines = content.split("\n")
    if not lines:
        return content

    first = lines[0].strip().lstrip("#").strip().strip("*").strip()
    normalized = title.strip().lower()

    if first.lower() == normalized or first.lower().startswith(f"{normalized}:"):
        remainder = "\n".join(lines[1:]).strip()
        if remainder:
            return remainder
    return content
