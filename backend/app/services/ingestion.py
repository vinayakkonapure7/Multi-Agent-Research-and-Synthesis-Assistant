"""
Ingestion pipeline for user-supplied sources.

Flow per source:
  file/url -> raw text extraction -> chunk -> embed -> store in Chroma (namespaced by project_id)

OCR strategy for images / scanned & handwritten PDFs (per user's choice: local-first, Vision fallback):
  1. Try Tesseract OCR locally (fast, free).
  2. If result is too short / low quality (common for handwritten notes, messy scans),
     fall back to GPT-4o-mini Vision, which handles handwriting far better.
"""
import io
import base64
import logging
from pathlib import Path

import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import httpx
import trafilatura
from docx import Document
try:
    import openpyxl
except Exception:  # pragma: no cover - optional until dependencies are installed
    openpyxl = None

from app.core.config import settings
from app.services.vector_store import get_collection
from app.services.llm_client import vision_extract_text

logger = logging.getLogger(__name__)

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150
MIN_OCR_CHARS = 25  # below this, treat local OCR as "failed" and try Vision fallback


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = text.strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap
        if start < 0 or end >= len(text):
            break
    return [c for c in chunks if c.strip()]


def recursive_character_chunks(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """
    RecursiveCharacterTextSplitter-style chunking. It prefers paragraph,
    newline, sentence, and word boundaries before falling back to characters.
    """
    text = text.strip()
    if not text:
        return []

    separators = ["\n\n", "\n", ". ", " ", ""]
    chunks = []

    def split_piece(piece: str, level: int = 0):
        piece = piece.strip()
        if not piece:
            return
        if len(piece) <= chunk_size:
            chunks.append(piece)
            return
        sep = separators[level] if level < len(separators) else ""
        if sep == "":
            for start in range(0, len(piece), max(1, chunk_size - overlap)):
                chunks.append(piece[start:start + chunk_size].strip())
            return

        current = ""
        for part in piece.split(sep):
            part = part.strip()
            candidate = f"{current}{sep}{part}".strip() if current else part
            if len(candidate) <= chunk_size:
                current = candidate
            else:
                if current:
                    chunks.append(current)
                if len(part) > chunk_size:
                    split_piece(part, level + 1)
                    current = ""
                else:
                    current = part
        if current:
            chunks.append(current)

    split_piece(text)
    if overlap <= 0:
        return [c for c in chunks if c.strip()]

    with_overlap = []
    previous_tail = ""
    for chunk in chunks:
        merged = f"{previous_tail} {chunk}".strip() if previous_tail else chunk
        with_overlap.append(merged[:chunk_size])
        previous_tail = chunk[-overlap:]
    return [c for c in with_overlap if c.strip()]


def _ocr_image_local(img: Image.Image) -> str:
    try:
        return pytesseract.image_to_string(img)
    except Exception as e:
        logger.warning(f"Local OCR failed: {e}")
        return ""


def _image_to_b64(img: Image.Image) -> str:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def extract_text_from_image(path: Path) -> str:
    """Local Tesseract first; fallback to GPT-4o-mini Vision for handwriting/poor scans."""
    img = Image.open(path).convert("RGB")
    local_text = _ocr_image_local(img)

    if len(local_text.strip()) >= MIN_OCR_CHARS:
        return local_text

    logger.info(f"Local OCR weak ({len(local_text.strip())} chars) for {path.name}; trying Vision fallback")
    b64 = _image_to_b64(img)
    vision_text = vision_extract_text(b64)
    return vision_text or local_text


def extract_text_from_pdf(path: Path) -> str:
    """
    Extracts embedded text per page. For pages with little/no extractable text
    (i.e. scanned/handwritten pages), rasterizes that page and runs the image OCR chain.
    """
    doc = fitz.open(path)
    full_text = []
    for page_index in range(len(doc)):
        page = doc[page_index]
        page_text = page.get_text().strip()

        if len(page_text) >= MIN_OCR_CHARS:
            full_text.append(page_text)
            continue

        # Likely scanned/handwritten page -> rasterize and OCR
        pix = page.get_pixmap(dpi=200)
        img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
        local_text = _ocr_image_local(img)

        if len(local_text.strip()) >= MIN_OCR_CHARS:
            full_text.append(local_text)
        else:
            b64 = _image_to_b64(img)
            vision_text = vision_extract_text(b64)
            full_text.append(vision_text or local_text)

    doc.close()
    return "\n\n".join(full_text)


def extract_text_from_txt(path: Path) -> str:
    return path.read_text(errors="ignore")


def extract_text_from_docx(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def extract_text_from_xlsx(path: Path) -> str:
    if openpyxl is None:
        raise RuntimeError("openpyxl is required to ingest .xlsx files. Install backend requirements again.")
    workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
    parts = []
    for sheet in workbook.worksheets:
        parts.append(f"Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if values:
                parts.append(" | ".join(values))
    workbook.close()
    return "\n".join(parts)


def extract_text_from_url(url: str) -> str:
    """Fetch and extract main article content from a URL, stripping nav/ads/boilerplate."""
    resp = httpx.get(url, timeout=20.0, follow_redirects=True,
                      headers={"User-Agent": "Mozilla/5.0 (research-report-builder)"})
    resp.raise_for_status()
    extracted = trafilatura.extract(resp.text, url=url, include_tables=True)
    return extracted or ""


def ingest_source(project_id: str, source_id: str, kind: str, name: str, file_path: str | None) -> int:
    """
    Extracts + chunks + embeds a single source. Returns chunk count.
    Raises on failure (caller updates Source.status accordingly).
    """
    if kind == "pdf":
        text = extract_text_from_pdf(Path(file_path))
    elif kind == "image":
        text = extract_text_from_image(Path(file_path))
    elif kind in {"text", "json", "csv"}:
        text = extract_text_from_txt(Path(file_path))
    elif kind == "docx":
        text = extract_text_from_docx(Path(file_path))
    elif kind == "xlsx":
        text = extract_text_from_xlsx(Path(file_path))
    elif kind == "url":
        text = extract_text_from_url(name)
    else:
        raise ValueError(f"Unknown source kind: {kind}")

    if not text.strip():
        raise ValueError("No extractable text found in source")

    chunks = recursive_character_chunks(text)
    if not chunks:
        raise ValueError("Text extracted but produced no usable chunks")

    collection = get_collection(project_id)
    ids = [f"{source_id}_{i}" for i in range(len(chunks))]
    metadatas = [{"source_id": source_id, "source_name": name, "kind": kind, "chunk_index": i}
                 for i in range(len(chunks))]
    collection.add(documents=chunks, ids=ids, metadatas=metadatas)

    return len(chunks)
